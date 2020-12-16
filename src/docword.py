import docx
import os
from docx.shared import Cm

# Variables 
titular_de_la_cuenta = [bd_cliente_estudio_de_ahorro_modelo["Titular_de_la_Cuenta"]
direccion_suministro = [bd_cliente_estudio_de_ahorro_modelo["Dirección_completa"]
cups = [bd_cliente_estudio_de_ahorro_modelo["Cups"]
nueva_comercializadora = the_one_comercializadora["Comercializadora"]
ahorro_percent = porcentage_de_ahorro[0]
ahorro_euros = ahorro[0]
records = [bd_cliente_estudio_de_ahorro_modelo["Precio_Energia_P1"], bd_cliente_estudio_de_ahorro_modelo["Precio_Energia_P2"], bd_cliente_estudio_de_ahorro_modelo["Precio_Energia_P3"]] # precio energia cliente
records_2 = [bd_cliente_estudio_de_ahorro_modelo["Precio_Potencia_P1"], bd_cliente_estudio_de_ahorro_modelo["Precio_Potencia_P2"], bd_cliente_estudio_de_ahorro_modelo["Precio_Potencia_P3"]] # precio potencia cliente
records_3 = [the_one_comercializadora["Precio_Energía_P1"],the_one_comercializadora["Precio_Energía_P2"],the_one_comercializadora["Precio_Energía_P3"]] # precio nuevo energia
records_4 = [the_one_comercializadora["Precio_Potencia_P1"],the_one_comercializadora["Precio_Potencia_P2"],the_one_comercializadora["Precio_Potencia_P3"]] # precio nuevo potencia 


# Pagina 1 Word Doc

doc =  docx.Document()

vivolt = doc.add_picture('../Data/Imagenes/VIVOLT logo 2.png', width=Cm(15))
vivolt.alignment = 1

primer_heading = doc.add_heading('Estudio de Ahorro', 0).bold=True

primer_subparag = doc.add_paragraph()
primer_subparag.add_run('Ahorra en tu factura de luz, de manera 100% gratuita y con energía 100% renovable').italic=True

segundo_heading = doc.add_heading('Nuestra Propuesta', 1)

#valor que insertar
parag1_propuesta = doc.add_paragraph()
parag1_propuesta.add_run('Estudio Preparado para: ')
parag1_propuesta.add_run(titular_de_la_cuenta)

#valor que insertar
parag2_propuesta = doc.add_paragraph()
parag2_propuesta.add_run('Dirección del punto de suministro: ')
parag2_propuesta.add_run(direccion_suministro)

#valor que insertar
parag3_propuesta = doc.add_paragraph()
parag3_propuesta.add_run('CUPS: ')
parag3_propuesta.add_run(cups)

#valor que insertar
parag9_propuesta = doc.add_paragraph()
parag9_propuesta.add_run('Este año, la mejor comercializadora para ti es: ')
parag9_propuesta.add_run(nueva_comercializadora).bold=True

#valor que insertar
parag4_propuesta = doc.add_paragraph()
parag4_propuesta.add_run('Ahorro : ')
parag4_propuesta.add_run(ahorro_percent).bold=True

#valor que insertar
parag5_propuesta = doc.add_paragraph()
parag5_propuesta.add_run('Ahorro anual : ')
parag5_propuesta.add_run(ahorro_euros).bold=True

#valor que insertar. 
parag6_propuesta = doc.add_paragraph()
parag6_propuesta.add_run('Emisiones de CO2/kg ahorradas : ')
parag6_propuesta.add_run('emisiones(multiplicar 1kwh por 0,385)').bold=True

#valor que insertar
parag7_propuesta = doc.add_paragraph()
parag7_propuesta.add_run('El ahorro de tus emisiones de CO2 equivale a plantar : ')
parag7_propuesta.add_run('arbol= 0.4 CO2/kg, dividir emisiones por arbol').bold=True
parag7_propuesta.add_run('arboles').bold=True

parag8_propuesta = doc.add_paragraph()
parag8_propuesta.add_run('Este ')
parag8_propuesta.add_run('estudio personalizado ').bold=True
parag8_propuesta.add_run('optimiza tus condiciones tarifarias, analizando tus patrones de consumo, contrastando tus condiciones actuales contra la oferta del mercado y así poder ajustarlas a las que mas te convengan.')

parag10_propuesta = doc.add_paragraph()
parag10_propuesta.alignment = 1
parag10_propuesta.paragraph_format.line_spacing = 8
parag10_propuesta.add_run('¡Apuntate a VIVOLT y nosotros nos ocupamos de siempre el mejor precio').bold=True

# Pagina 2 Word Doc
doc.add_page_break()

tercer_heading = doc.add_heading('Nuestra Propuesta Detallada', 1).bold=True

cuarto_heading = doc.add_heading('Antes tenías:', 2).bold=True

quinto_heading = doc.add_heading('Precio Energía:', 3).bold=True

records = [
    ['0.3535435', "0.34353554", '0.545456']
           ]
# cambiar al real

antes_energia_table = doc.add_table(rows=1, cols=3)
antes_energia_table.style= 'Medium Shading 1'
hdr_Cells = antes_energia_table.rows[0].cells
hdr_Cells[0].text = 'P1'
hdr_Cells[1].text = 'P2'
hdr_Cells[2].text = 'P3'
for p1, p2, p3 in records:
    row_Cells = antes_energia_table.add_row().cells
    row_Cells[0].text = p1
    row_Cells[1].text = p2
    row_Cells[2].text = p3

sexto_heading = doc.add_heading('Precio Potencia:', 3).bold=True

records_2 = [
    ['0.8535435', "0.74353554", '0.645456']
           ]
# cambiar al real

antes_potencia_table = doc.add_table(rows=1, cols=3)
antes_potencia_table.style= 'Medium Shading 1'
hdr_Cells = antes_potencia_table.rows[0].cells
hdr_Cells[0].text = 'P1'
hdr_Cells[1].text = 'P2'
hdr_Cells[2].text = 'P3'
for p1, p2, p3 in records_2:
    row_Cells = antes_potencia_table.add_row().cells
    row_Cells[0].text = p1
    row_Cells[1].text = p2
    row_Cells[2].text = p3

septimo_heading = doc.add_heading('Ahora puedes tener:', 2).bold=True

octavo_heading = doc.add_heading('Precio Energía:', 3).bold=True

records_3 = [
    ['0.8535435', "0.74353554", '0.645456']
           ]
# cambiar al real

nueva_energia_table = doc.add_table(rows=1, cols=3)
nueva_energia_table.style= 'Light List Accent 6'
hdr_Cells = nueva_energia_table.rows[0].cells
hdr_Cells[0].text = 'P1'
hdr_Cells[1].text = 'P2'
hdr_Cells[2].text = 'P3'
for p1, p2, p3 in records_3:
    row_Cells = nueva_energia_table.add_row().cells
    row_Cells[0].text = p1
    row_Cells[1].text = p2
    row_Cells[2].text = p3

noveno_heading = doc.add_heading('Precio Potencia:', 3).bold=True

records_4 = [
    ['0.8535435', "0.74353554", '0.645456']
           ]
# cambiar al real

nueva_potencia_table = doc.add_table(rows=1, cols=3)
nueva_potencia_table.style= 'Light List Accent 6'
hdr_Cells = nueva_potencia_table.rows[0].cells
hdr_Cells[0].text = 'P1'
hdr_Cells[1].text = 'P2'
hdr_Cells[2].text = 'P3'
for p1, p2, p3 in records_4:
    row_Cells = nueva_potencia_table.add_row().cells
    row_Cells[0].text = p1
    row_Cells[1].text = p2
    row_Cells[2].text = p3

decimo_heading = doc.add_heading('Tu consumo se distribuye de la siguiente manera:', 1).bold=True

pie_chart = doc.add_picture('../Data/Imagenes/pie_chart_consumo_cliente.png', width=Cm(10))
pie_chart.alignment = 1

parag_piechart_propuesta = doc.add_paragraph()
parag_piechart_propuesta.add_run('P1, P2 y P3 representan franjas horarias en las que consumes energía. Estas franjas horarias varían segun invierno o verano, aunque podríamos decir que:')

parag_p1_propuesta = doc.add_paragraph()
parag_p1_propuesta.add_run('P1: ').bold=True
parag_p1_propuesta.add_run('Tarde')

parag_p2_propuesta = doc.add_paragraph()
parag_p2_propuesta.add_run('P2: ').bold=True
parag_p2_propuesta.add_run('Mañana')

parag_p3_propuesta = doc.add_paragraph()
parag_p3_propuesta.add_run('P3: ').bold=True
parag_p3_propuesta.add_run('Noche')